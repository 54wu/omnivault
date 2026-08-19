# -*- coding: utf-8 -*-
"""
一键离线转文本：扫描一个文件夹里的 Word / PDF / 图片 / 文本，
全部转成统一的 Markdown(.md) 输出。

用法:
    python convert.py <材料目录> [输出目录] [--md|--txt]

说明:
    - .docx            -> python-docx 提取文本, 表格转为 md 表格
    - .pdf             -> 优先抽文本层; 若为扫描件(无文本层)则逐页转图片 OCR
    - .jpg/.jpeg/.png/.bmp/.webp/.tif/.tiff -> RapidOCR 识别
    - .txt/.md/.csv    -> 原样复制
    - 默认输出 .md; 传 --txt 则输出纯文本 .txt
    - 全程离线, 不联网。
"""

import sys
from pathlib import Path

# ---------- 支持的格式 ----------
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tif", ".tiff"}
TEXT_EXTS = {".txt", ".md", ".csv", ".log"}
ALL_EXTS = IMAGE_EXTS | TEXT_EXTS | {".docx", ".pdf"}


def _table_to_md(table) -> str:
    """把 docx 表格转成 md 表格。"""
    rows = []
    for row in table.rows:
        rows.append([c.text.replace("\n", " ").strip() for c in row.cells])
    if not rows:
        return ""
    out = []
    header = rows[0]
    out.append("| " + " | ".join(header) + " |")
    out.append("| " + " | ".join("---" for _ in header) + " |")
    for r in rows[1:]:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    if doc.tables:
        lines.append("")
        for t in doc.tables:
            lines.append(_table_to_md(t))
            lines.append("")
    return "\n".join(lines)


def extract_pdf(path: Path, ocr) -> str:
    import fitz  # PyMuPDF
    parts = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            text = page.get_text("text").strip()
            if text:
                parts.append(text)
            else:
                # 无文本层 -> 扫描件, 转图片 OCR
                pix = page.get_pixmap(dpi=200)
                import io
                from PIL import Image
                img_bytes = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_bytes))
                result, _ = ocr(img)
                page_text = "\n".join(line[1] for line in (result or []))
                parts.append(page_text)
    return "\n".join(p for p in parts if p)


def extract_image(path: Path, ocr) -> str:
    result, _ = ocr(str(path))
    if not result:
        return ""
    return "\n".join(line[1] for line in result)


def convert_file(path: Path, ocr) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".docx":
            return extract_docx(path)
        if ext == ".pdf":
            return extract_pdf(path, ocr)
        if ext in IMAGE_EXTS:
            return extract_image(path, ocr)
        if ext in TEXT_EXTS:
            return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        return f"\n[解析失败 {path.name}: {e}]\n"


def _iter_files(src_dir: Path, out_dir: Path):
    """递归扫描 src_dir 下所有可处理文件, 跳过 out_dir 内部。"""
    out_resolved = out_dir.resolve()
    for p in src_dir.rglob("*"):
        if not p.is_file() or p.suffix.lower() not in ALL_EXTS:
            continue
        if out_resolved in p.resolve().parents:
            continue
        yield p


def main():
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        sys.exit(1)

    fmt = "md"
    if "--txt" in args:
        fmt = "txt"
    args = [a for a in args if a not in ("--md", "--txt")]

    src_dir = Path(args[0]).resolve()
    out_dir = Path(args[1]).resolve() if len(args) > 1 else src_dir / "_output"

    if not src_dir.is_dir():
        print(f"错误: 找不到目录 {src_dir}")
        sys.exit(1)
    out_dir.mkdir(parents=True, exist_ok=True)

    ext = ".md" if fmt == "md" else ".txt"

    # 惰性初始化 OCR(卡内会稍微慢, 只有用到图片/扫描件才加载)
    ocr = None

    # 跳过输出目录本身, 避免把已生成的输出当输入递归处理
    files = list(_iter_files(src_dir, out_dir))
    print(f"共发现 {len(files)} 个可处理文件 (输出格式: {ext})\n")

    for i, f in enumerate(files, 1):
        if f.suffix.lower() in IMAGE_EXTS or f.suffix.lower() == ".pdf":
            if ocr is None:
                from rapidocr_onnxruntime import RapidOCR
                print("首次遇到图片/扫描件, 正在加载 OCR 引擎...")
                ocr = RapidOCR()
        text = convert_file(f, ocr)
        rel = f.relative_to(src_dir)
        out_path = out_dir / rel.with_suffix(ext)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(text, encoding="utf-8")
        lines = text.count("\n") + 1 if text.strip() else 0
        print(f"[{i}/{len(files)}] {f.name} -> {lines} 行  {rel}")

    print(f"\n完成! 输出目录: {out_dir}")


if __name__ == "__main__":
    main()